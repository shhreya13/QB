# endsem_qb_final.py
import streamlit as st
import re, random
from io import BytesIO
from copy import deepcopy
from collections import Counter, defaultdict
from docx import Document

st.set_page_config(page_title="EndSem QB Generator (Final)", layout="wide")

# -----------------------
# Regex helpers
# -----------------------
BLOOM_RE = re.compile(r'K\s*([1-6])', re.I)
CO_RE = re.compile(r'CO\s*[_:]?\s*(\d+)', re.I)
UNIT_RE = re.compile(r'Unit\s*[-:]?\s*(\d+)', re.I)
NUM_PREFIX_RE = re.compile(r'^\s*(\d+)\s*[\.\)]')  # matches "1." or "1)"

# -----------------------
# Low-level copy helper
# -----------------------
def _copy_element(elem):
    return deepcopy(elem)

def replace_cell_with_cell(target_cell, src_cell):
    """
    Replace the XML of target_cell with src_cell's XML (deep copy).
    Preserves inline OMML and drawing nodes (images).
    """
    t_tc = target_cell._tc
    s_tc = src_cell._tc
    # clear target children
    for child in list(t_tc):
        t_tc.remove(child)
    # append deep-copied children from source
    for child in list(s_tc):
        t_tc.append(_copy_element(child))

# -----------------------
# Parsing question bank (tables)
# -----------------------
def extract_questions_from_bank_docx(uploaded_file):
    """
    Expects question bank in tables. Each question row should contain Bloom (K#) and optionally CO and Unit.
    Returns list of q dicts:
      {'id':int, 'unit':int|None, 'co':int|None, 'bloom':int, 'cell':<Cell>, 'text':str}
    """
    doc = Document(uploaded_file)
    questions = []
    qid = 0
    for table in doc.tables:
        for row in table.rows:
            texts = [c.text.strip() for c in row.cells]
            if all(not t for t in texts):
                continue
            joined = " | ".join(texts)
            bm = BLOOM_RE.search(joined)
            if not bm:
                continue  # skip rows without bloom tag
            um = UNIT_RE.search(joined)
            cm = CO_RE.search(joined)
            main_idx = max(range(len(texts)), key=lambda i: len(texts[i]))
            main_cell = row.cells[main_idx]
            qid += 1
            q = {
                'id': qid,
                'unit': int(um.group(1)) if um else None,
                'co': int(cm.group(1)) if cm else None,
                'bloom': int(bm.group(1)),
                'cell': main_cell,
                'text': main_cell.text.strip()
            }
            questions.append(q)
    return questions

# -----------------------
# Parse template: find numbered slots and OR markers in tables
# -----------------------
def parse_template_slots(template_file):
    doc = Document(template_file)
    slots = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if not txt:
                    continue
                up = txt.upper()
                # OR marker detection
                if up == '(OR)' or up == 'OR' or up == '( OR )':
                    slots.append({
                        'table_index': ti, 'row_index': ri, 'cell_index': ci,
                        'slot_num': None, 'is_or': True, 'unit_in_template': None
                    })
                    continue
                # numeric slot detection e.g., "11." or " 12 )"
                mnum = NUM_PREFIX_RE.match(txt)
                if mnum:
                    slot_num = int(mnum.group(1))
                    unitm = UNIT_RE.search(txt)
                    unit_in_template = int(unitm.group(1)) if unitm else None
                    slots.append({
                        'table_index': ti, 'row_index': ri, 'cell_index': ci,
                        'slot_num': slot_num, 'is_or': False, 'unit_in_template': unit_in_template
                    })
                    continue
                # detect standalone Unit placeholder
                unitm = UNIT_RE.search(txt)
                if unitm:
                    slots.append({
                        'table_index': ti, 'row_index': ri, 'cell_index': ci,
                        'slot_num': None, 'is_or': False, 'unit_in_template': int(unitm.group(1))
                    })
    return slots

# -----------------------
# Bloom range by slot number
# -----------------------
def allowed_blooms_for_slot_num(slot_num):
    if slot_num is None:
        return [1,2,3,4,5,6]
    if 1 <= slot_num <= 10:
        return [1,2,3]
    if 11 <= slot_num <= 20:
        return [4,5]
    if 21 <= slot_num <= 22:
        return [6]
    return [1,2,3,4,5,6]

# -----------------------
# Map slots and find OR pairs
# -----------------------
def map_slots_to_template_positions(slots):
    entries = []
    for s in slots:
        if s['slot_num'] is not None:
            entries.append({
                'slot_num': s['slot_num'],
                'table_index': s['table_index'],
                'row_index': s['row_index'],
                'cell_index': s['cell_index'],
                'unit_in_template': s.get('unit_in_template')
            })
    entries.sort(key=lambda x: x['slot_num'])
    return entries

def find_or_pairs_from_slots(slots):
    pairs = []
    for idx, s in enumerate(slots):
        if s['is_or']:
            prev_slot = None
            next_slot = None
            j = idx - 1
            while j >= 0:
                if slots[j]['slot_num'] is not None:
                    prev_slot = slots[j]
                    break
                j -= 1
            k = idx + 1
            while k < len(slots):
                if slots[k]['slot_num'] is not None:
                    next_slot = slots[k]
                    break
                k += 1
            if prev_slot and next_slot:
                pairs.append((prev_slot, next_slot))
    return pairs

# -----------------------
# Select questions for slots
# -----------------------
def select_questions_for_slots(entries, or_pairs, questions, allow_reuse=False):
    selected = {}
    used_ids = set()

    by_unit_bloom = defaultdict(lambda: defaultdict(list))
    by_bloom = defaultdict(list)
    for q in questions:
        by_bloom[q['bloom']].append(q)
        if q['unit'] is not None:
            by_unit_bloom[q['unit']][q['bloom']].append(q)

    slotnum_to_entry = {e['slot_num']: e for e in entries}

    # Handle OR pairs first
    for left_slot, right_slot in or_pairs:
        left_num, right_num = left_slot['slot_num'], right_slot['slot_num']
        left_entry = slotnum_to_entry.get(left_num)
        right_entry = slotnum_to_entry.get(right_num)
        if not left_entry or not right_entry:
            continue
        left_allowed = allowed_blooms_for_slot_num(left_num)
        right_allowed = allowed_blooms_for_slot_num(right_num)
        preferred_unit = left_entry.get('unit_in_template') or right_entry.get('unit_in_template')
        candidate_units = [preferred_unit] if preferred_unit else []
        candidate_units += sorted(set(q['unit'] for q in questions if q['unit'] is not None and q['unit']!=preferred_unit))
        chosen_pair = None
        for unit in candidate_units:
            if unit is None:
                continue
            left_cands = [q for b in left_allowed for q in by_unit_bloom.get(unit, {}).get(b, [])]
            right_cands = [q for b in right_allowed for q in by_unit_bloom.get(unit, {}).get(b, [])]
            if not left_cands or not right_cands:
                continue
            found = None
            for lc in left_cands:
                for rc in right_cands:
                    if lc['id'] != rc['id'] or allow_reuse:
                        found = (lc, rc)
                        break
                if found:
                    break
            if found:
                chosen_pair = found
                break
        if not chosen_pair:
            left_pool = [q for b in left_allowed for q in by_bloom.get(b, [])]
            right_pool = [q for b in right_allowed for q in by_bloom.get(b, [])]
            if not allow_reuse:
                left_pool = [q for q in left_pool if q['id'] not in used_ids]
                right_pool = [q for q in right_pool if q['id'] not in used_ids]
            if left_pool and right_pool:
                lc = random.choice(left_pool)
                if lc in right_pool and not allow_reuse:
                    right_pool = [q for q in right_pool if q['id'] != lc['id']]
                rc = random.choice(right_pool) if right_pool else random.choice([q for q in by_bloom.get(right_allowed[0], [])])
                chosen_pair = (lc, rc)
            else:
                all_cands = [q for q in questions if q['bloom'] in (left_allowed + right_allowed)]
                if all_cands:
                    a, b = random.sample(all_cands, 2) if len(all_cands) >= 2 else (all_cands[0], all_cands[0])
                    chosen_pair = (a, b)
        left_coord = (left_entry['table_index'], left_entry['row_index'], left_entry['cell_index'])
        right_coord = (right_entry['table_index'], right_entry['row_index'], right_entry['cell_index'])
        lc, rc = chosen_pair
        selected[left_coord] = lc
        selected[right_coord] = rc
        if not allow_reuse:
            used_ids.update([lc['id'], rc['id']])

    # handle remaining slots
    for e in entries:
        coord = (e['table_index'], e['row_index'], e['cell_index'])
        if coord in selected:
            continue
        slot_num = e['slot_num']
        blooms = allowed_blooms_for_slot_num(slot_num)
        cand_pool = []
        if e.get('unit_in_template'):
            u = e['unit_in_template']
            for b in blooms:
                cand_pool += by_unit_bloom.get(u, {}).get(b, [])
        if not cand_pool:
            for b in blooms:
                cand_pool += by_bloom.get(b, [])
        if not cand_pool:
            cand_pool = questions[:]
        if not allow_reuse:
            cand_pool = [q for q in cand_pool if q['id'] not in used_ids]
            if not cand_pool:
                cand_pool = [q for q in questions if q['id'] not in used_ids] or questions[:]
        chosen = random.choice(cand_pool)
        selected[coord] = chosen
        if not allow_reuse:
            used_ids.add(chosen['id'])

    return selected

# -----------------------
# Assemble doc by replacing only Unit Placeholder (2nd column)
# -----------------------
def assemble_generated_doc(template_file, bank_file, selected_map):
    tmpl = Document(template_file)
    bank = Document(bank_file)

    for (ti, ri, ci), q in selected_map.items():
        try:
            t_table = tmpl.tables[ti]
            t_row = t_table.rows[ri]
        except Exception:
            continue

        # Replace ONLY the second column (Unit Placeholder)
        if len(t_row.cells) >= 2:
            target_cell = t_row.cells[1]
            replace_cell_with_cell(target_cell, q['cell'])

        # Fill CO_ and K_ placeholders in the row if present
        for idx, c in enumerate(t_row.cells):
            txt = c.text.strip()
            up = txt.upper()
            if 'CO' in up and ('_' in up or up.strip() == 'CO' or up.strip().startswith('CO')):
                if q.get('co'):
                    t_row.cells[idx].text = f"CO {q['co']}"
                else:
                    t_row.cells[idx].text = "CO"
            if re.search(r'\bK[_]?\b', up) or up.strip().startswith('K_') or up.strip().startswith('K'):
                if q.get('bloom'):
                    t_row.cells[idx].text = f"K{q['bloom']}"
                else:
                    t_row.cells[idx].text = "K_"

    buf = BytesIO()
    tmpl.save(buf)
    buf.seek(0)
    return buf

# -----------------------
# Streamlit UI
# -----------------------
st.title("End-sem Question Paper Generator — Final")

st.markdown("""
**Instructions**
- Upload your **Template DOCX** (table-based) and **Question Bank DOCX** (table-based).
- Template must have numbered question slots like `1.`, `2.` ... `22.` matching the numbering scheme.
- The app will populate:
  - Q1–Q10 → K1/K2/K3
  - Q11–Q20 → K4/K5
  - Q21–Q22 → K6
- `(OR)` markers between numbered rows will be handled: the two choices are filled with questions from the **same unit** (preferred).
- Only the second column (Unit Placeholder) will be replaced with the question text; S.No, CO, and K columns remain unchanged.
""")

template_file = st.file_uploader("Upload Template (.docx)", type=["docx"])
bank_file = st.file_uploader("Upload Question Bank (.docx)", type=["docx"])

n_sets = st.number_input("Number of sets to generate", min_value=1, max_value=20, value=1, step=1)
allow_reuse = st.checkbox("Allow reuse of the same question within a set", value=False)

if st.button("Generate"):
    if not template_file or not bank_file:
        st.error("Please upload both Template and Question Bank files.")
    else:
        template_file.seek(0)
        bank_file.seek(0)

        with st.spinner("Parsing question bank..."):
            questions = extract_questions_from_bank_docx(bank_file)
        st.success(f"Found {len(questions)} candidate questions in bank.")
        if len(questions) == 0:
            st.error("No questions with K-levels found in the uploaded bank. Make sure each question row contains 'K1'..'K6'.")
        else:
            with st.spinner("Parsing template slots and OR markers..."):
                template_file.seek(0)
                raw_slots = parse_template_slots(template_file)
                entries = map_slots_to_template_positions(raw_slots)
                or_pairs = find_or_pairs_from_slots(raw_slots)
            st.write(f"Template discovered {len(entries)} numeric slots and {len(or_pairs)} OR pairs (if any).")

            all_buffers = []
            usage_counter = Counter()
            for s in range(n_sets):
                template_file.seek(0)
                bank_file.seek(0)
                selected_map = select_questions_for_slots(entries, or_pairs, questions, allow_reuse)
                for q in selected_map.values():
                    usage_counter[q['id']] += 1
                buf = assemble_generated_doc(template_file, bank_file, selected_map)
                all_buffers.append((buf, selected_map))

            total_placed = sum(usage_counter.values())
            unique_used = len(usage_counter)
            repetition_pct = (1 - unique_used / total_placed) * 100 if total_placed > 0 else 0.0

            st.subheader("Dashboard")
            st.write(f"Total placed questions across all sets: {total_placed}")
            st.write(f"Unique questions used: {unique_used}")
            st.write(f"Repetition percentage: {repetition_pct:.2f}%")

            unit_counts = Counter(q['unit'] for q in questions if q['unit'] is not None)
            st.subheader("Unit weightage in bank")
            total_qs = sum(unit_counts.values())
            for u, c in sorted(unit_counts.items()):
                st.write(f"Unit {u}: {c} questions -> { (c / total_qs * 100) if total_qs else 0:.1f}%")

            st.subheader("Generated sets (download & preview)")
            for i, (buf, sel_map) in enumerate(all_buffers, start=1):
                st.download_button(label=f"Download Set {i}", data=buf, file_name=f"QuestionPaper_Set{i}.docx")
                with st.expander(f"Preview Set {i} (first 12 items)"):
                    preview_list = []
                    for coord, q in sel_map.items():
                        ti, ri, ci = coord
                        tmpl = Document(template_file)
                        try:
                            cell_txt = tmpl.tables[ti].rows[ri].cells[ci].text.strip()
                        except Exception:
                            cell_txt = ""
                        sn = None
                        mm = NUM_PREFIX_RE.match(cell_txt)
                        if mm:
                            sn = int(mm.group(1))
                        else:
                            sn = ""
                        preview_list.append((sn, q['unit'], q['co'], q['bloom'], q['text'][:250].replace('\n',' ')))
                    preview_list_sorted = sorted(preview_list, key=lambda x: (x[0] if isinstance(x[0], int) else 9999))
                    for p in preview_list_sorted[:12]:
                        st.markdown(f"**Slot {p[0]}** | Unit: {p[1]} | CO: {p[2]} | K{p[3]} — {p[4]}")

            st.info("Notes: Only the second column (Unit Placeholder) is replaced. OMML and images are preserved. S.No, CO, K columns remain unchanged.")
